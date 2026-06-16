#!/usr/bin/env python3
"""Generate Word document for AI recognition prompt annotations."""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[3]
MD_PATH = ROOT / "docs" / "prompt-annotation.md"
OUT_DIR = ROOT / "docs" / "sop" / "output"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "AttendanceAgent_提示词注释说明.docx"

VERSION = "1.0.0"
DOC_DATE = date.today().isoformat()


def set_doc_font(doc: Document, east_asia: str = "PingFang SC"):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI 考勤识别提示词注释说明")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("AttendanceAgent Prompt Annotation Guide")
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"\nVersion {VERSION}  |  {DOC_DATE}\n")
    meta.add_run("源文件：docs/prompt-annotation.md\n")
    meta.add_run("配置源：base-config/prompts.md\n")
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int):
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False, indent: float = 0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold


def add_code_block(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def parse_table_rows(block: str) -> list[list[str]]:
    rows = []
    for line in block.strip().splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[-:\s|]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows


def add_markdown_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < cols:
                table.rows[i].cells[j].text = cell
                if i == 0:
                    for p in table.rows[i].cells[j].paragraphs:
                        for r in p.runs:
                            r.bold = True
    doc.add_paragraph()


def render_markdown(doc: Document, content: str):
    lines = content.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith("#### "):
            add_heading(doc, line[5:].strip(), 4)
        elif line.strip() == "---":
            doc.add_paragraph()
        elif line.strip().startswith("|"):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, parse_table_rows("\n".join(table_lines)))
            continue
        elif line.strip().startswith("> "):
            p = doc.add_paragraph(line.strip()[2:])
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        elif line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", line.strip()):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line.strip()), style="List Number")
        elif line.strip():
            text = line.strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"`([^`]+)`", r"\1", text)
            add_paragraph(doc, text)

        i += 1


def main():
    if not MD_PATH.exists():
        raise SystemExit(f"Markdown source not found: {MD_PATH}")

    doc = Document()
    set_doc_font(doc)
    add_cover(doc)
    render_markdown(doc, MD_PATH.read_text(encoding="utf-8"))
    doc.save(OUT_PATH)
    print(f"Generated: {OUT_PATH}")


if __name__ == "__main__":
    main()
