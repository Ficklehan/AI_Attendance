#!/usr/bin/env python3
"""Generate multilingual Attendance Agent SOP Word document with embedded screenshots."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[3]
SHOT = Path(__file__).resolve().parent.parent / "screenshots"
OUT = Path(__file__).resolve().parent.parent / "output"
OUT.mkdir(parents=True, exist_ok=True)

VERSION = "1.0.0"
DOC_DATE = date.today().isoformat()


def set_doc_font(doc: Document, east_asia: str = "PingFang SC"):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_cover(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"\nVersion {VERSION}  |  {DOC_DATE}\n")
    meta.add_run("AI Attendance Assistant / AI考勤智能助手\n")
    meta.add_run("Languages: 简体中文 · English · Français · Deutsch\n")
    doc.add_page_break()


def add_image(doc: Document, filename: str, caption: str, width_cm: float = 15.5):
    path = SHOT / filename
    if path.exists():
        doc.add_picture(str(path), width=Cm(width_cm))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    else:
        doc.add_paragraph(f"[Screenshot pending: {filename}] — {caption}")
    doc.add_paragraph()


def add_steps(doc: Document, steps: list[tuple[str, str]]):
    for i, (action, expected) in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {action}", style="List Number")
        p = doc.add_paragraph(f"   Expected: {expected}")
        p.paragraph_format.left_indent = Cm(0.8)
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def add_troubleshooting(doc: Document, rows: list[tuple[str, str, str]]):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Symptom", "Cause", "Resolution"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, (sym, cause, fix) in enumerate(rows, 1):
        row = table.rows[ri].cells
        row[0].text = sym
        row[1].text = cause
        row[2].text = fix
    doc.add_paragraph()


# ─── Language content ───────────────────────────────────────────────────────

LANG_SECTIONS = {
    "zh": {
        "label": "第一部分 · 简体中文",
        "toc_title": "目录（中文）",
        "sections": [
            ("1. 系统概述", [
                "AI考勤智能助手用于将纸质/拍照考勤表通过 AI 视觉识别自动结构化，经人工核对后写入飞书多维表格。",
                "系统包含：PC Web 管理端、飞书小程序现场端、Spring Boot 后端 API。",
                "支持多国家/地区配置（中、法、德、美、荷、波、意、西、捷等），各国可使用独立的 AI 提示词与飞书字段映射。",
            ]),
            ("2. 角色与权限", [
                "普通用户：上传识别、核对确认、查看本人任务、切换工作国家、导出本人数据。",
                "管理员：上述全部 + 查看全员任务、AI/飞书配置、用户管理、审计日志。",
                "重要：在「设置 → 用户管理」绑定飞书 open_id 后，PC 账号与小程序登录共享同一 user_id 下的任务数据。",
            ]),
            ("3. 任务状态说明", [
                "processing（识别中）→ processed（待核对）→ confirmed（已完成，同步飞书）",
                "failed（识别失败，可重新上传）；cancelled（已作废）",
                "统计数字统一来自 GET /attendance/api/tasks/summary，勿用列表分页条数代替全库统计。",
            ]),
            ("4. PC 端：上传与识别", [
                ("打开浏览器访问系统（默认 http://localhost:5175），使用账号登录。", "01-home-upload.png", "图 4-1：首页上传区"),
                ("在首页确认「当前工作国家」；如需更改，点击「更改配置」进入设置。", None, None),
                ("点击或拖拽上传考勤图片（JPG/PNG，支持多张，单张≤10MB）。", None, None),
                ("点击「开始识别」，等待 AI 流式解析完成（通常 30–90 秒）。", None, None),
                ("在结果页核对识别行，修正异常字段后保存。", None, None),
            ]),
            ("5. PC 端：任务列表与核对", [
                ("进入「任务列表」，查看识别中/待核对/已完成/失败数量。", "02-task-list.png", "图 5-1：任务列表"),
                ("点击「去核对 (N)」或筛选「待核对」状态，打开任务详情。", None, None),
                ("在任务编辑页逐行检查姓名、日期、到达/离开时间、休息分钟等字段。", None, None),
                ("确认无误后点击「确认提交」，系统将数据写入飞书多维表格。", None, None),
                ("若同步失败，在任务详情点击「重试飞书同步」。", None, None),
            ]),
            ("6. PC 端：考勤记录与导出", [
                ("进入「考勤记录」，按员工维度查看所有已识别明细。", "03-task-records.png", "图 6-1：考勤记录"),
                ("使用搜索/高级搜索筛选工号、姓名、日期等字段。", None, None),
                ("点击「导出 Excel」创建异步导出任务，完成后下载文件。", None, None),
            ]),
            ("7. 飞书小程序：现场采集", [
                ("在飞书工作台打开「AI考勤助手」小程序，使用飞书账号授权登录。", "08-mini-login.png", "图 7-1：小程序登录"),
                ("在「识别」页选择工作国家，拍照或从相册选择考勤表。", "09-mini-workbench.png", "图 7-2：识别工作台"),
                ("进入识别中页面等待 AI 完成（见进度与步骤提示）。", "10-mini-processing.png", "图 7-3：识别进行中"),
                ("在结果页修正异常行，确认后提交；数据同步至飞书。", None, None),
                ("在「任务」页查看各状态任务，支持左滑删除。", None, None),
            ]),
            ("8. 管理员：系统配置", [
                ("设置 → AI 识别：维护各国识别提示词与质量规则。", None, None),
                ("设置 → 飞书：配置多维表格 app_token、table_id 及字段映射。", None, None),
                ("设置 → 用户管理：创建账号、分配角色、绑定飞书 open_id。", None, None),
                ("设置 → 审计日志：查看登录、确认、配置变更等操作记录。", None, None),
            ]),
        ],
        "steps_upload": [
            ("登录 PC 端并进入首页", "显示当前工作国家与上传区域"),
            ("上传考勤图片并点击「开始识别」", "任务状态变为「识别中」"),
            ("等待识别完成", "状态变为「待核对」，可查看识别结果"),
            ("核对并确认提交", "状态变为「已完成」，飞书同步成功"),
        ],
        "troubleshoot": [
            ("识别结果为空或乱码", "图片模糊、国家配置不匹配", "重拍清晰照片；确认工作国家与表格格式一致"),
            ("飞书同步失败", "Bitable 权限或字段映射错误", "检查 feishu 配置；任务详情点「重试同步」"),
            ("PC 与小程序任务不一致", "未绑定同一飞书 open_id", "管理员在用户管理中绑定 open_id"),
            ("待核对数量与列表不符", "使用了分页计数", "以 summary 接口统计为准，刷新页面"),
        ],
    },
    "en": {
        "label": "Part 2 · English",
        "toc_title": "Table of Contents (English)",
        "sections": [
            ("1. System Overview", [
                "AI Attendance Assistant digitizes paper attendance sheets via MiMo Vision AI, human review, and Feishu Bitable sync.",
                "Clients: Vue 3 Web admin, Feishu mini-program for field staff, Spring Boot API backend.",
                "Multi-country configs (CN, FR, DE, US, NL, PL, IT, ES, CZ, etc.) with per-country AI prompts and Feishu field mappings.",
            ]),
            ("2. Roles & Permissions", [
                "User: upload, recognize, review, confirm own tasks, switch work country, export own data.",
                "Admin: all above + all-users task view, AI/Feishu config, user management, audit logs.",
                "Bind Feishu open_id under Settings → Users so PC and mini-program share the same task data.",
            ]),
            ("3. Task Status Lifecycle", [
                "processing → processed (pending review) → confirmed (synced to Feishu)",
                "failed (retry upload); cancelled",
                "Counts must come from GET /attendance/api/tasks/summary — never use paginated list length as totals.",
            ]),
            ("4. PC: Upload & Recognition", [
                ("Open the web app and sign in (default http://localhost:5175).", "01-home-upload.png", "Fig 4-1: Home upload"),
                ("Verify working country on Home; change via Settings if needed.", None, None),
                ("Upload JPG/PNG images (multi-file, max 10MB each).", None, None),
                ("Click Start Recognition and wait for streaming AI parse (~30–90s).", None, None),
                ("Review rows on the result view; fix anomalies before saving.", None, None),
            ]),
            ("5. PC: Task List & Review", [
                ("Open Task List; check counts for Processing / Pending review / Completed / Failed.", "02-task-list.png", "Fig 5-1: Task list"),
                ("Use Go to review (N) or filter by Pending review.", None, None),
                ("Open task detail; verify name, date, arrival/departure, break minutes.", None, None),
                ("Click Confirm to write records to Feishu Bitable.", None, None),
                ("On sync failure, use Retry Feishu sync on the task detail page.", None, None),
            ]),
            ("6. PC: Records & Export", [
                ("Open Attendance Records for employee-level detail.", "03-task-records.png", "Fig 6-1: Records"),
                ("Use search / advanced search filters.", None, None),
                ("Export Excel creates an async job; download when ready.", None, None),
            ]),
            ("7. Feishu Mini-Program: Field Workflow", [
                ("Open the mini-app in Feishu; authorize with Feishu account.", "08-mini-login.png", "Fig 7-1: Mini login"),
                ("On Recognize tab, set country; capture or pick photos.", "09-mini-workbench.png", "Fig 7-2: Workbench"),
                ("Wait on Recognizing screen until AI completes.", "10-mini-processing.png", "Fig 7-3: Processing"),
                ("Review results, fix anomalies, confirm and sync to Feishu.", None, None),
                ("Tasks tab: filter by status; swipe to delete.", None, None),
            ]),
            ("8. Admin: Configuration", [
                ("Settings → AI: country prompts and quality rules.", None, None),
                ("Settings → Feishu: Bitable tokens and field mapping.", None, None),
                ("Settings → Users: accounts, roles, Feishu open_id binding.", None, None),
                ("Settings → Audit: operation history.", None, None),
            ]),
        ],
        "steps_upload": [
            ("Sign in to PC and open Home", "Working country and upload zone visible"),
            ("Upload images and start recognition", "Task status = Processing"),
            ("Wait for recognition to finish", "Status = Pending review"),
            ("Review and confirm", "Status = Completed; Feishu sync OK"),
        ],
        "troubleshoot": [
            ("Empty or garbled recognition", "Blurry image or wrong country config", "Retake photo; match working country to sheet format"),
            ("Feishu sync failed", "Bitable permission or field mapping", "Check Feishu config; retry sync on task"),
            ("PC vs mini-program data mismatch", "Feishu open_id not bound", "Admin binds open_id in User management"),
            ("Review count mismatch", "Used page size as total", "Use summary API; refresh page"),
        ],
    },
    "fr": {
        "label": "Partie 3 · Français",
        "toc_title": "Table des matières (Français)",
        "sections": [
            ("1. Vue d'ensemble", [
                "L'assistant de pointage IA numérise les feuilles papier via MiMo Vision, revue humaine et synchronisation Feishu Bitable.",
                "Clients : Web Vue 3, mini-programme Feishu, API Spring Boot.",
                "Configurations multi-pays avec invites IA et mapping de champs Feishu par pays.",
            ]),
            ("2. Rôles", [
                "Utilisateur : téléverser, reconnaître, vérifier, confirmer ses tâches, changer le pays de travail.",
                "Administrateur : tout + tâches de tous les utilisateurs, config IA/Feishu, utilisateurs, audit.",
                "Lier open_id Feishu pour unifier les données PC et mini-programme.",
            ]),
            ("3. Cycle de statut", [
                "processing → processed (à vérifier) → confirmed (synchronisé Feishu)",
                "failed ; cancelled",
                "Comptages via GET /attendance/api/tasks/summary uniquement.",
            ]),
            ("4. PC : Téléversement", [
                ("Ouvrir l'application Web et se connecter.", "01-home-upload.png", "Fig 4-1 : Accueil"),
                ("Vérifier le pays de travail.", None, None),
                ("Téléverser JPG/PNG (multi-fichiers, max 10 Mo).", None, None),
                ("Lancer la reconnaissance IA.", None, None),
                ("Vérifier les lignes extraites.", None, None),
            ]),
            ("5. PC : Liste des tâches", [
                ("Ouvrir Liste des tâches.", "02-task-list.png", "Fig 5-1"),
                ("Filtrer « À vérifier » ou cliquer Aller vérifier.", None, None),
                ("Corriger les champs puis Confirmer.", None, None),
                ("Réessayer la sync Feishu si échec.", None, None),
            ]),
            ("6. PC : Enregistrements", [
                ("Ouvrir Enregistrements de pointage.", "03-task-records.png", "Fig 6-1"),
                ("Recherche et export Excel.", None, None),
            ]),
            ("7. Mini-programme Feishu", [
                ("Connexion Feishu.", "08-mini-login.png", "Fig 7-1"),
                ("Onglet Reconnaissance : photo ou album.", "09-mini-workbench.png", "Fig 7-2"),
                ("Attendre la fin de l'IA.", "10-mini-processing.png", "Fig 7-3"),
                ("Vérifier et confirmer.", None, None),
            ]),
            ("8. Admin", [
                ("Paramètres → IA, Feishu, Utilisateurs, Audit.", None, None),
            ]),
        ],
        "steps_upload": [
            ("Connexion PC, page d'accueil", "Pays de travail affiché"),
            ("Téléverser et lancer la reconnaissance", "Statut = En cours"),
            ("Attendre la fin", "Statut = À vérifier"),
            ("Confirmer", "Statut = Terminé, sync Feishu"),
        ],
        "troubleshoot": [
            ("Reconnaissance vide", "Image floue ou mauvais pays", "Reprendre la photo"),
            ("Échec sync Feishu", "Permissions ou mapping", "Vérifier config Feishu"),
            ("Données PC / mini différentes", "open_id non lié", "Lier open_id"),
        ],
    },
    "de": {
        "label": "Teil 4 · Deutsch",
        "toc_title": "Inhaltsverzeichnis (Deutsch)",
        "sections": [
            ("1. Systemübersicht", [
                "Der KI-Anwesenheitsassistent digitalisiert Papier-Stundenzettel per MiMo Vision, manueller Prüfung und Feishu-Bitable-Sync.",
                "Clients: Vue-3-Web, Feishu-Miniprogramm, Spring-Boot-API.",
                "Mehrländer-Konfiguration mit länderspezifischen KI-Prompts und Feishu-Feldzuordnungen.",
            ]),
            ("2. Rollen", [
                "Benutzer: Upload, Erkennung, Prüfung, Bestätigung eigener Aufgaben, Arbeitsland wechseln.",
                "Administrator: alles + alle Benutzeraufgaben, KI-/Feishu-Konfiguration, Benutzerverwaltung, Audit.",
                "Feishu open_id binden für gemeinsame Daten zwischen PC und Mini-Programm.",
            ]),
            ("3. Aufgabenstatus", [
                "processing → processed (zu prüfen) → confirmed (mit Feishu synchronisiert)",
                "failed; cancelled",
                "Zählungen nur über GET /attendance/api/tasks/summary.",
            ]),
            ("4. PC: Upload", [
                ("Web-App öffnen und anmelden.", "01-home-upload.png", "Abb. 4-1"),
                ("Arbeitsland prüfen.", None, None),
                ("JPG/PNG hochladen.", None, None),
                ("Erkennung starten.", None, None),
                ("Ergebnisse prüfen.", None, None),
            ]),
            ("5. PC: Aufgabenliste", [
                ("Aufgabenliste öffnen.", "02-task-list.png", "Abb. 5-1"),
                ("Filter „Zu prüfen“ oder Prüfen starten.", None, None),
                ("Bestätigen → Feishu-Sync.", None, None),
            ]),
            ("6. PC: Datensätze", [
                ("Anwesenheitsdatensätze öffnen.", "03-task-records.png", "Abb. 6-1"),
                ("Suche und Excel-Export.", None, None),
            ]),
            ("7. Feishu-Miniprogramm", [
                ("Feishu-Anmeldung.", "08-mini-login.png", "Abb. 7-1"),
                ("Erkennungs-Tab: Foto/Album.", "09-mini-workbench.png", "Abb. 7-2"),
                ("KI-Verarbeitung abwarten.", "10-mini-processing.png", "Abb. 7-3"),
            ]),
            ("8. Admin", [
                ("Einstellungen → KI, Feishu, Benutzer, Audit.", None, None),
            ]),
        ],
        "steps_upload": [
            ("PC-Anmeldung, Startseite", "Arbeitsland sichtbar"),
            ("Bilder hochladen, Erkennung starten", "Status = In Bearbeitung"),
            ("Warten", "Status = Zu prüfen"),
            ("Bestätigen", "Status = Abgeschlossen"),
        ],
        "troubleshoot": [
            ("Leere Erkennung", "Unscharfes Bild", "Neues Foto"),
            ("Feishu-Sync fehlgeschlagen", "Berechtigung/Mapping", "Config prüfen"),
            ("PC/Mini unterschiedlich", "open_id fehlt", "open_id binden"),
        ],
    },
}


def render_language_section(doc: Document, lang: dict):
    h = doc.add_heading(lang["label"], level=1)
    h.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    doc.add_heading(lang["toc_title"], level=2)
    for sec_title, items in lang["sections"]:
        doc.add_paragraph(sec_title, style="List Bullet")
    doc.add_paragraph()

    for sec_title, items in lang["sections"]:
        doc.add_heading(sec_title, level=2)
        for item in items:
            if isinstance(item, str):
                doc.add_paragraph(item)
            elif isinstance(item, tuple):
                text, img, cap = item
                doc.add_paragraph(text)
                if img and cap:
                    add_image(doc, img, cap)
        doc.add_paragraph()

    doc.add_heading("Quick workflow", level=2)
    add_steps(doc, lang["steps_upload"])

    doc.add_heading("Troubleshooting", level=2)
    add_troubleshooting(doc, lang["troubleshoot"])
    doc.add_page_break()


def build_document() -> Path:
    doc = Document()
    set_doc_font(doc)

    add_cover(
        doc,
        "AI Attendance Assistant\nStandard Operating Procedure",
        "Multilingual System SOP · 多语言系统操作手册",
    )

    # Document control
    doc.add_heading("Document Control / 文档控制", level=1)
    ctrl = doc.add_table(rows=6, cols=2)
    ctrl.style = "Table Grid"
    rows = [
        ("Document ID", "AA-SOP-001"),
        ("Version", VERSION),
        ("Date", DOC_DATE),
        ("Owner", "Product Management"),
        ("Audience", "Field HR, Back-office, System Admins"),
        ("Related skill", ".cursor/skills/attendance-pm-agent/"),
    ]
    for i, (k, v) in enumerate(rows):
        ctrl.rows[i].cells[0].text = k
        ctrl.rows[i].cells[1].text = v
    doc.add_paragraph()
    doc.add_page_break()

    # Architecture diagram (text)
    doc.add_heading("System Architecture / 系统架构", level=1)
    doc.add_paragraph(
        "Feishu Mini-Program → API (Spring Boot) → MiMo Vision AI\n"
        "                              ↓\n"
        "                         MySQL tasks\n"
        "                              ↓\n"
        "                      Feishu Bitable sync"
    )
    doc.add_page_break()

    for key in ("zh", "en", "fr", "de"):
        render_language_section(doc, LANG_SECTIONS[key])

    # PM Agent appendix
    doc.add_heading("Appendix: Product Manager Agent", level=1)
    doc.add_paragraph(
        "Invoke the attendance-pm-agent skill in Cursor to generate PRDs, update this SOP, "
        "or analyze feature impact. Regenerate this document:"
    )
    doc.add_paragraph("python docs/sop/scripts/capture_screenshots.py", style="Intense Quote")
    doc.add_paragraph("python docs/sop/scripts/generate_sop_docx.py", style="Intense Quote")

    out_path = OUT / "AttendanceAgent_SOP_Multilingual.docx"
    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    path = build_document()
    print(f"Generated: {path}")
