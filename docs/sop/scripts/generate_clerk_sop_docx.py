#!/usr/bin/env python3
"""Generate warehouse-clerk SOP Word document (zh-CN) with embedded screenshots."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[3]
SHOT = Path(__file__).resolve().parent.parent / "screenshots"
OUT = Path(__file__).resolve().parent.parent / "output"
OUT.mkdir(parents=True, exist_ok=True)

VERSION = "1.1.0"
DOC_DATE = date.today().strftime("%Y年%m月%d日")


def set_doc_font(doc: Document, east_asia: str = "PingFang SC"):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def shade_cell(cell, fill: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI 考勤智能助手\n仓库文员操作手册")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Standard Operating Procedure · 仓库现场文员版")
    s.font.size = Pt(13)
    s.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"\n文档编号：AA-SOP-CLERK-001\n版本：{VERSION}  |  日期：{DOC_DATE}\n")
    meta.add_run("适用角色：仓库文员 / 现场考勤录入人员\n")
    meta.add_run("涵盖端：PC 网页端 + 飞书小程序\n")
    doc.add_page_break()


def add_image(doc: Document, filename: str, caption: str, width_cm: float = 15.0):
    path = SHOT / filename
    if path.exists():
        doc.add_picture(str(path), width=Cm(width_cm))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    else:
        doc.add_paragraph(f"【截图待补充：{filename}】{caption}")
    doc.add_paragraph()


def add_callout(doc: Document, title: str, lines: list[str], fill: str = "FFF9EC"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    t = p.add_run(f"⚠ {title}\n")
    t.bold = True
    t.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    for line in lines:
        p.add_run(f"• {line}\n")
    doc.add_paragraph()


def add_step(doc: Document, n: int, action: str, note: str | None = None):
    p = doc.add_paragraph(style="List Number")
    p.add_run(action)
    if note:
        q = doc.add_paragraph(f"　注意：{note}")
        q.paragraph_format.left_indent = Cm(0.6)
        for r in q.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def build_document() -> Path:
    doc = Document()
    set_doc_font(doc)
    add_cover(doc)

    # ── 阅读导引 ──
    doc.add_heading("如何使用本手册", level=1)
    doc.add_paragraph(
        "本手册按「先总后分、先安全后操作」编排：请先阅读第 1 章注意事项，再按章节完成日常操作。"
        "每节包含：操作目的 → 分步指引（配图）→ 本节要点。遇到异常请直接查阅第 7 章常见问题。"
    )
    doc.add_paragraph()

    # ── 快速流程 ──
    doc.add_heading("30 秒速览：每日标准流程", level=1)
    flow = doc.add_table(rows=5, cols=3)
    flow.style = "Table Grid"
    headers = ["步骤", "做什么", "完成标志"]
    for i, h in enumerate(headers):
        flow.rows[0].cells[i].text = h
        shade_cell(flow.rows[0].cells[i], "EFF6FF")
    rows = [
        ("1", "登录系统（PC 或飞书）", "进入首页 / 识别页"),
        ("2", "确认工作国家/地区 → 拍照或上传考勤表", "任务状态「识别中」"),
        ("3", "等待 AI 识别 → 打开「待核对」任务", "可看到逐行识别结果"),
        ("4", "修正异常字段 → 确认提交", "状态「已完成」，飞书同步成功"),
    ]
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            flow.rows[ri].cells[ci].text = val
    doc.add_paragraph()
    doc.add_paragraph("查看历史明细：PC 端进入「考勤记录」，按工号/姓名/日期筛选。")
    doc.add_page_break()

    # ── 第1章 注意事项 ──
    doc.add_heading("第 1 章  操作前必读（重要）", level=1)
    add_callout(doc, "以下情况会导致识别错误或无法提交，请务必遵守", [
        "拍照前确认「当前工作国家/地区」与纸质考勤表所属国家一致（如法国表选 FR）。",
        "照片需平整、光线充足、文字清晰；避免反光、裁切到表格边缘、手指遮挡。",
        "单张图片不超过 10MB，格式 JPG/PNG；多张图片会合并为同一识别任务。",
        "「NO」是线下纸质表的序号，不是系统工号；确认后系统自动分配系统工号，请勿手工改序号逻辑。",
        "必填字段（姓名、日期、到达/离开时间、仓库、中介等）缺失时无法确认，须补全或标记删除。",
        "已确认的任务不可删除；如需重录须联系管理员处理。",
        "飞书同步失败时，数据已保存在系统内，可在任务详情点击「重试飞书同步」。",
    ])
    doc.add_heading("任务状态说明", level=2)
    status = doc.add_table(rows=6, cols=2)
    status.style = "Table Grid"
    status.rows[0].cells[0].text = "状态"
    status.rows[0].cells[1].text = "含义与文员操作"
    shade_cell(status.rows[0].cells[0], "EFF6FF")
    shade_cell(status.rows[0].cells[1], "EFF6FF")
    for i, (s, m) in enumerate([
        ("识别中", "AI 正在解析，请等待，可离开页面后台继续"),
        ("待核对", "识别完成，必须人工检查后再确认"),
        ("已完成", "已确认并写入飞书，一般无需再改"),
        ("识别失败", "图片质量问题或网络异常，请重拍重传"),
        ("已作废", "任务已取消，不计入考勤"),
    ], 1):
        status.rows[i].cells[0].text = s
        status.rows[i].cells[1].text = m
    doc.add_page_break()

    # ── 第2章 登录 ──
    doc.add_heading("第 2 章  登录系统", level=1)
    doc.add_paragraph("目的：进入系统并确保使用本人账号，保证任务数据归属正确。")

    doc.add_heading("2.1  PC 网页端登录", level=2)
    add_step(doc, 1, "打开浏览器，访问考勤系统地址（由 IT 提供，开发环境示例：http://localhost:5175/clockai/）")
    add_step(doc, 2, "输入管理员分配的用户名和密码，点击「登录」", "首次登录建议修改密码")
    add_step(doc, 3, "也可点击「飞书登录」——需管理员已在后台绑定您的飞书账号", "未绑定时飞书登录后看不到历史任务")
    add_image(doc, "00-login.png", "图 2-1  PC 登录页")
    add_step(doc, 4, "登录成功后自动进入首页，右上角可查看当前用户与工作国家")

    doc.add_heading("2.2  飞书小程序登录", level=2)
    add_step(doc, 1, "在飞书工作台搜索并打开「AI 考勤助手」小程序")
    add_step(doc, 2, "点击「使用飞书登录」，按提示完成授权", "须使用公司飞书账号")
    add_step(doc, 3, "首次登录后，在 PC「用户管理」绑定飞书 ID 可与网页端共享任务", "如仅使用小程序可跳过 PC 绑定，但数据仅限本机飞书账号")
    add_image(doc, "08-mini-login.png", "图 2-2  飞书小程序登录页", width_cm=8.5)
    doc.add_page_break()

    # ── 第3章 拍照识别 ──
    doc.add_heading("第 3 章  拍照识别考勤表", level=1)
    doc.add_paragraph("目的：将纸质考勤表拍照/上传，由 AI 自动提取姓名、日期、到离岗时间等字段。")

    doc.add_heading("3.1  PC 端上传识别", level=2)
    add_step(doc, 1, "登录后进入「首页」", "确认页面上方显示的「当前工作国家」正确")
    add_image(doc, "01-home-upload.png", "图 3-1  PC 首页上传区")
    add_step(doc, 2, "点击上传区域或拖拽图片到虚线框内，可选择多张", "建议按页码顺序上传")
    add_step(doc, 3, "点击「开始识别」", "上传后任务进入「识别中」，大表首张约 1～2 分钟")
    add_step(doc, 4, "识别完成后首页/任务列表出现「待核对」提示，点击「去核对」进入详情", "识别过程中可处理其他工作，不必一直等待")

    doc.add_heading("3.2  飞书小程序拍照识别", level=2)
    add_step(doc, 1, "打开小程序底部「识别」页，检查顶部国家/语言标签", "点击「更改」可切换工作国家")
    add_image(doc, "09-mini-workbench.png", "图 3-2  小程序识别工作台", width_cm=8.5)
    add_step(doc, 2, "选择「拍照」现场拍摄，或「从相册选择」已有照片", "拍摄时尽量垂直对准表格，四边完整入镜")
    add_step(doc, 3, "多张图片会进入待识别队列，点击底部主按钮「开始识别」", "队列可清空后重新选择")
    add_step(doc, 4, "进入识别中页面，等待进度条与「条已解析」数字增长", "可点「返回首页」或「查看任务」，识别在后台继续")
    add_image(doc, "10-mini-processing.png", "图 3-3  小程序识别进行中", width_cm=8.5)
    doc.add_page_break()

    # ── 第4章 核对确认 ──
    doc.add_heading("第 4 章  核对与确认提交", level=1)
    doc.add_paragraph(
        "目的：人工校验 AI 识别结果，修正错误后正式确认。此步骤不可跳过——确认后数据将写入飞书多维表格。"
    )
    add_callout(doc, "核对重点（按优先级）", [
        "日期格式是否正确（与纸质表一致）",
        "到达 / 离开时间是否合理（离开早于到达可能是夜班跨日，留意系统标记）",
        "标有「看不清」「???」的字段必须补全或删除该行",
        "重名人员留意「重名疑似」提示，按主管指示处理",
        "未出勤行保留 SmartMark「未出勤」，不要强行填写时间",
    ], fill="EFF6FF")

    doc.add_heading("4.1  PC 端核对", level=2)
    add_step(doc, 1, "进入「任务列表」，查看「待核对」数量，或筛选状态为「待核对」", "列表数字以系统统计为准，勿用单页条数估算")
    add_image(doc, "02-task-list.png", "图 4-1  任务列表与状态筛选")
    add_step(doc, 2, "点击任务编号进入「任务详情」", "页面上方显示本任务工作地区")
    add_image(doc, "04-task-edit.png", "图 4-2  任务详情与行级编辑")
    add_step(doc, 3, "逐行检查表格，直接在单元格内修改错误字段；红色/黄色标记行优先处理")
    add_step(doc, 4, "若有必填缺失，底部会提示无法提交；点击「查看详情」定位到具体行号")
    add_step(doc, 5, "全部无误后点击「确认提交」", "提交后状态变为「已完成」，系统自动同步飞书")
    add_step(doc, 6, "若提示飞书同步失败，点击「重试飞书同步」", "数据已保存，重试不会丢失")

    doc.add_heading("4.2  飞书小程序核对", level=2)
    add_step(doc, 1, "在「任务」页切换至「待核对」，点击任务卡片进入结果页")
    add_image(doc, "13-mini-tasks.png", "图 4-3  小程序任务列表", width_cm=8.5)
    add_step(doc, 2, "查看顶部统计：总条数 / 有效 / 待修正；黄色提示栏表示仍有必填缺失")
    add_image(doc, "12-mini-result.png", "图 4-4  小程序核对结果页", width_cm=8.5)
    add_step(doc, 3, "点击「编辑此行」修正单条记录；标红字段为必填或格式错误")
    add_step(doc, 4, "全部修正后点击「确认提交（同步飞书）」", "提交前再次确认工作国家与表头一致")
    doc.add_page_break()

    # ── 第5章 考勤记录 ──
    doc.add_heading("第 5 章  查看考勤明细记录", level=1)
    doc.add_paragraph("目的：按员工/日期查询已确认的历史考勤行，用于对账与抽查。")

    doc.add_heading("5.1  PC 端「考勤记录」", level=2)
    add_step(doc, 1, "左侧菜单点击「考勤记录」", "仅显示已确认任务中的有效行")
    add_image(doc, "03-task-records.png", "图 5-1  考勤记录列表")
    add_step(doc, 2, "使用顶部状态下拉筛选「已完成」记录（默认已排除已删除行）")
    add_step(doc, 3, "输入关键词快速搜索工号、姓名、仓库等；点击「高级搜索」按字段精确筛选")
    add_step(doc, 4, "点击「导出 Excel」可创建异步导出任务，完成后在导出中心下载", "大批量导出请避开业务高峰")
    add_step(doc, 5, "点击列头可排序；通过列设置隐藏/冻结常用列", None)

    doc.add_heading("5.2  小程序查看任务", level=2)
    add_step(doc, 1, "「任务」页切换「已完成」查看已确认任务")
    add_step(doc, 2, "点击任务可回看识别结果（已确认任务为只读，不可再改）", "需修改须联系管理员")
    doc.add_page_break()

    # ── 第6章 术语 ──
    doc.add_heading("第 6 章  术语速查", level=1)
    glossary = doc.add_table(rows=1, cols=2)
    glossary.style = "Table Grid"
    glossary.rows[0].cells[0].text = "术语"
    glossary.rows[0].cells[1].text = "说明"
    shade_cell(glossary.rows[0].cells[0], "EFF6FF")
    shade_cell(glossary.rows[0].cells[1], "EFF6FF")
    terms = [
        ("工作国家/地区", "决定 AI 识别规则与飞书写入目标表，须与纸质表国家一致"),
        ("NO / 线下序号", "纸质考勤表左侧手写序号，识别后保留原值"),
        ("系统工号", "确认后系统自动分配（如 FR00001），用于员工维度汇总"),
        ("待核对", "AI 识别完成、等待人工确认的状态"),
        ("SmartMark / 标记", "系统对模糊、手写、未出勤等行的自动标签"),
        ("飞书同步", "确认后将数据写入飞书多维表格的过程"),
    ]
    for term, desc in terms:
        row = glossary.add_row().cells
        row[0].text = term
        row[1].text = desc
    doc.add_paragraph()
    doc.add_page_break()

    # ── 第7章 FAQ ──
    doc.add_heading("第 7 章  常见问题", level=1)
    faq = doc.add_table(rows=1, cols=2)
    faq.style = "Table Grid"
    faq.rows[0].cells[0].text = "现象"
    faq.rows[0].cells[1].text = "处理办法"
    shade_cell(faq.rows[0].cells[0], "EFF6FF")
    shade_cell(faq.rows[0].cells[1], "EFF6FF")
    for q, a in [
        ("识别结果大量为空或乱码", "检查照片是否模糊；确认工作国家是否选对；重拍后重新上传"),
        ("确认按钮灰色/无法提交", "查看底部必填校验提示，补全标红字段或删除无效行"),
        ("飞书同步失败", "在任务详情点「重试飞书同步」；持续失败联系管理员检查飞书配置"),
        ("PC 与小程序任务不一致", "确认两端登录为同一账号（PC 需绑定飞书 open_id）"),
        ("找不到某天的记录", "确认任务是否已「确认提交」；在考勤记录中扩大日期筛选范围"),
        ("识别一直停在「识别中」", "等待 2～3 分钟；仍无进展刷新页面或到任务列表查看；联系 IT"),
    ]:
        row = faq.add_row().cells
        row[0].text = q
        row[1].text = a

    doc.add_paragraph()
    doc.add_paragraph("—— 文档结束 ——").alignment = WD_ALIGN_PARAGRAPH.CENTER

    out = OUT / "仓库文员操作手册_AI考勤智能助手.docx"
    doc.save(str(out))
    return out


if __name__ == "__main__":
    path = build_document()
    print(f"Generated: {path}")
